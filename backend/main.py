from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import google.generativeai as genai
import io
try:
    import fitz  # PyMuPDF (optional on some hosts)
except Exception:
    fitz = None
import pdfplumber
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import os
import tempfile
import re
from typing import Dict, List
import json
import asyncio
import anyio
import subprocess
import shutil
from fastapi import BackgroundTasks

from pathlib import Path
import base64
import requests
try:
    import PyPDF2  # optional for page count
except Exception:
    PyPDF2 = None

app = FastAPI(title="HireMe Maker")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global variable to store API key for session
gemini_api_key = None

@app.get("/")
async def root():
    return {"message": "HireMe Maker API"}

@app.post("/set-api-key")
async def set_api_key(api_key: str = Form(...)):
    """Store Gemini API key (minimal validation to avoid false negatives)."""
    global gemini_api_key
    key = api_key.strip()
    if not key:
        raise HTTPException(status_code=400, detail="API key is required")
    
    # Simply configure the API key without validation
    # This allows us to accept any key format and test it only when needed
    try:
        genai.configure(api_key=key)
        gemini_api_key = key
        return {"status": "success", "message": "API key stored"}
    except HTTPException:
        raise
    except Exception as e:
        # Only catch configuration errors, not validation errors
        raise HTTPException(status_code=400, detail=f"Failed to configure API: {str(e)}")

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from uploaded resume file"""
    if filename.lower().endswith('.pdf'):
        # Try with PyMuPDF first if available
        if fitz is not None:
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                text = "".join(page.get_text() for page in doc)
                doc.close()
                return text
            except Exception:
                pass
        # Fallback to pdfplumber (pure-Python)
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
            return text
    
    elif filename.lower().endswith(('.doc', '.docx')):
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    elif filename.lower().endswith('.txt'):
        return file_content.decode('utf-8')
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported file format")

def calculate_ats_score(resume_text: str, job_description: str) -> Dict:
    """Calculate ATS score based on keyword matching"""
    # Extract keywords from job description
    jd_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', job_description.lower()))
    resume_words = set(re.findall(r'\b[a-zA-Z]{3,}\b', resume_text.lower()))
    
    # Common stop words to exclude (deduplicated)
    stop_words = {
        'the','and','for','are','but','not','you','all','can','had','her','was','one','our','out','day','get','has','him','his','how','its',
        'may','new','old','see','two','who','did','man','way','she','use','your','said','each','which','their','time','will','about','would',
        'there','could','other','after','first','well','water','been','call','oil','sit','find','long','down','come','made','part'
    }
    
    jd_keywords = jd_words - stop_words
    resume_keywords = resume_words - stop_words
    
    # Find matching and missing keywords
    matching_keywords = jd_keywords.intersection(resume_keywords)
    missing_keywords = jd_keywords - resume_keywords
    
    # Calculate score
    if len(jd_keywords) > 0:
        score = (len(matching_keywords) / len(jd_keywords)) * 100
    else:
        score = 0
    
    return {
        "score": round(score, 2),
        "matching_keywords": list(matching_keywords)[:20],  # Limit to top 20
        "missing_keywords": list(missing_keywords)[:20],   # Limit to top 20
        "total_jd_keywords": len(jd_keywords),
        "matched_count": len(matching_keywords)
    }

@app.post("/analyze")
async def analyze_resume(
    resume: UploadFile = File(...),
    job_description: str = Form(...)
):
    """Analyze resume against job description and return ATS score"""
    try:
        # Extract text from resume
        resume_content = await resume.read()
        resume_text = extract_text_from_file(resume_content, resume.filename)
        
        # Calculate ATS score
        ats_analysis = calculate_ats_score(resume_text, job_description)
        
        return {
            "status": "success",
            "resume_text": resume_text,
            "ats_score": ats_analysis["score"],
            "matching_keywords": ats_analysis["matching_keywords"],
            "missing_keywords": ats_analysis["missing_keywords"],
            "analysis": ats_analysis
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

def escape_latex(text: str) -> str:
    if text is None:
        return ""
    # Escape LaTeX special characters
    replacements = {
        "\\": r"\textbackslash{}",
        "{": r"\{",
        "}": r"\}",
        "#": r"\#",
        "%": r"\%",
        "&": r"\&",
        "$": r"\$",
        "_": r"\_",
        "^": r"\^{}",
        "~": r"\~{}",
    }
    escaped = ""
    for ch in text:
        escaped += replacements.get(ch, ch)
    return escaped

def list_to_itemize(items: List[str]) -> str:
    if not items:
        return ""
    lines = ["\\begin{itemize}"]
    for it in items:
        if not it:
            continue
        lines.append(f"  \\item {escape_latex(str(it))}")
    lines.append("\\end{itemize}")
    return "\n".join(lines)

def render_latex_from_data(template_str: str, data: Dict) -> str:
    # Expected data keys
    name = escape_latex(data.get("name", ""))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    links = escape_latex(data.get("links", ""))

    summary = escape_latex(data.get("summary", ""))

    # Skills can be list[str] or dict[str, list[str]]
    skills_section = ""
    skills = data.get("skills")
    if isinstance(skills, dict):
        parts = []
        for group, items in skills.items():
            line = f"\\textbf{{{escape_latex(str(group))}}}: " + ", ".join(escape_latex(str(x)) for x in (items or []))
            parts.append(line)
        skills_section = "\\\n".join(parts)
    elif isinstance(skills, list):
        skills_section = ", ".join(escape_latex(str(x)) for x in skills)
    else:
        skills_section = escape_latex(str(skills) if skills else "")

    # Experience: list of {company, role, location, start, end, bullets:[]}
    exp_tex_parts = []
    for exp in (data.get("experience") or []):
        role = escape_latex(exp.get("role", ""))
        company = escape_latex(exp.get("company", ""))
        location = escape_latex(exp.get("location", ""))
        start = escape_latex(exp.get("start", ""))
        end = escape_latex(exp.get("end", ""))
        header = f"\\textbf{{{role}}} — {company}\\hfill {start}--{end}\\\n{location}\\\n"
        bullets = list_to_itemize(exp.get("bullets") or [])
        exp_tex_parts.append(header + bullets)
    experience_section = "\n\n".join(exp_tex_parts)

    # Projects: list of {name, url, description, bullets:[]}
    proj_parts = []
    for pr in (data.get("projects") or []):
        name = escape_latex(pr.get("name", ""))
        url = escape_latex(pr.get("url", ""))
        desc = escape_latex(pr.get("description", ""))
        header = f"\\textbf{{{name}}}\\hfill {url}\\\n{desc}\\\n"
        bullets = list_to_itemize(pr.get("bullets") or [])
        proj_parts.append(header + bullets)
    projects_section = "\n\n".join(proj_parts)

    # Education: list of {institution, degree, location, start, end, details}
    edu_parts = []
    for ed in (data.get("education") or []):
        inst = escape_latex(ed.get("institution", ""))
        degree = escape_latex(ed.get("degree", ""))
        location_e = escape_latex(ed.get("location", ""))
        start_e = escape_latex(ed.get("start", ""))
        end_e = escape_latex(ed.get("end", ""))
        details = escape_latex(ed.get("details", ""))
        edu_parts.append(f"\\textbf{{{degree}}}, {inst} — {location_e}\\hfill {start_e}--{end_e}\\\n{details}")
    education_section = "\n\n".join(edu_parts)

    # Certifications: list[str]
    certifications_section = list_to_itemize(data.get("certifications") or [])

    # Replace placeholders
    tex = template_str
    replacements = {
        "{{NAME}}": name,
        "{{EMAIL}}": email,
        "{{PHONE}}": phone,
        "{{LINKS}}": links,
        "{{SUMMARY}}": summary,
        "{{SKILLS}}": skills_section,
        "{{EXPERIENCE}}": experience_section,
        "{{PROJECTS}}": projects_section,
        "{{EDUCATION}}": education_section,
        "{{CERTIFICATIONS}}": certifications_section,
    }
    for k, v in replacements.items():
        tex = tex.replace(k, v or "")
    return tex

def find_latex_compiler() -> str:
    # Prefer latexmk, then pdflatex
    for cmd in ["latexmk", "pdflatex"]:
        if shutil.which(cmd):
            return cmd
    return ""

async def compile_tex_to_pdf_bytes(tex_source: str) -> bytes:
    workdir = Path(tempfile.mkdtemp(prefix="resume_tex_"))
    tex_file = workdir / "output.tex"
    tex_file.write_text(tex_source, encoding="utf-8")
    compiler = find_latex_compiler()
    if not compiler:
        raise HTTPException(status_code=500, detail="No LaTeX compiler found. Install MiKTeX/TeX Live and ensure pdflatex/latexmk in PATH.")
    if compiler == "latexmk":
        cmd = ["latexmk", "-pdf", "-interaction=nonstopmode", "-halt-on-error", tex_file.name]
        runs = 1
    else:
        cmd = ["pdflatex", "-interaction=nonstopmode", tex_file.name]
        runs = 2
    for _ in range(runs):
        proc = await asyncio.create_subprocess_exec(*cmd, cwd=str(workdir), stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT)
        stdout, _ = await proc.communicate()
        if proc.returncode != 0:
            log = stdout.decode(errors="ignore")
            shutil.rmtree(workdir, ignore_errors=True)
            raise HTTPException(status_code=500, detail=f"LaTeX compilation failed.\n{log[-4000:]}")
    pdf_path = workdir / "output.pdf"
    if not pdf_path.exists():
        candidate = workdir / f"{tex_file.stem}.pdf"
        if not candidate.exists():
            shutil.rmtree(workdir, ignore_errors=True)
            raise HTTPException(status_code=500, detail="PDF not generated by LaTeX compiler.")
        pdf_path = candidate
    data = pdf_path.read_bytes()
    shutil.rmtree(workdir, ignore_errors=True)
    return data

def render_cover_letter_from_data(template_str: str, data: Dict) -> str:
    name = escape_latex(data.get("name", ""))
    email = escape_latex(data.get("email", ""))
    phone = escape_latex(data.get("phone", ""))
    links = escape_latex(data.get("links", ""))
    company = escape_latex(data.get("company", ""))
    hiring_manager = escape_latex(data.get("hiring_manager", "Hiring Manager"))
    job_title = escape_latex(data.get("job_title", ""))
    opening = escape_latex(data.get("opening", ""))
    skills_fit = escape_latex(data.get("skills_fit", ""))
    conclusion = escape_latex(data.get("conclusion", ""))
    replacements = {
        "{{NAME}}": name,
        "{{EMAIL}}": email,
        "{{PHONE}}": phone,
        "{{LINKS}}": links,
        "{{COMPANY}}": company,
        "{{HIRING_MANAGER}}": hiring_manager,
        "{{JOB_TITLE}}": job_title,
        "{{OPENING}}": opening,
        "{{SKILLS_FIT}}": skills_fit,
        "{{CONCLUSION}}": conclusion,
    }
    tex = template_str
    for k, v in replacements.items():
        tex = tex.replace(k, v or "")
    return tex

def build_resume_prompt(resume_text: str, job_description: str) -> str:
    return (
        "Return ONLY valid JSON. Keys: name, email, phone, links, summary, "
        "skills (dict or list), experience (list of {role, company, location, start, end, bullets}), "
        "projects (list of {name, url, description, bullets}), education (list of {institution, degree, location, start, end, details}), "
        "certifications (list of strings). Keep content concise to fit a single A4 page.\n\n"
        f"Original Resume:\n{resume_text}\n\nJob Description:\n{job_description}"
    )

def build_cover_letter_prompt(resume_text: str, job_description: str, job_title: str, company: str) -> str:
    return (
        "Return ONLY valid JSON with keys: opening, skills_fit, conclusion. "
        "Keep to one page letter when rendered. Use ATS-friendly keywords.\n\n"
        f"Resume:\n{resume_text}\n\nJob Description:\n{job_description}\n\nRole: {job_title}\nCompany: {company}"
    )

@app.post("/tailor_resume")
async def tailor_resume_latex(payload: Dict):
    """Generate structured resume via Gemini (if data not provided), render LaTeX, compile to PDF and return base64 + latex."""
    global gemini_api_key
    data = payload.get("resume_data") or payload.get("data")
    username = payload.get("username") or (data.get("name") if isinstance(data, dict) else "tailored")
    if data is None:
        # Expect resume_text and job_description
        resume_text = payload.get("resume_text", "")
        job_description = payload.get("job_description", "")
        if not gemini_api_key:
            raise HTTPException(status_code=400, detail="API key not set")
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = build_resume_prompt(resume_text, job_description)
        resp = model.generate_content(prompt)
        try:
            data = json.loads(resp.text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gemini returned non-JSON: {str(e)}")

    # Render LaTeX
    base_dir = Path(__file__).parent
    template_path = base_dir / "templates" / "resume.tex"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="LaTeX template not found at backend/templates/resume.tex")
    template_str = template_path.read_text(encoding="utf-8")
    tex_source = render_latex_from_data(template_str, data)
    pdf_bytes = await compile_tex_to_pdf_bytes(tex_source)
    b64 = base64.b64encode(pdf_bytes).decode('ascii')
    safe_username = re.sub(r"[^A-Za-z0-9_-]+", "_", username or "tailored").strip("_") or "tailored"
    return {
        "status": "success",
        "filename": f"{safe_username}_resume.pdf",
        "latex": tex_source,
        "pdf_base64": b64,
        "data": data,
    }

@app.post("/generate_cover_letter")
async def generate_cover_letter(payload: Dict):
    """Generate tailored cover letter from resume/JD or structured paragraphs, return base64 PDF and LaTeX."""
    global gemini_api_key
    # Inputs
    name = payload.get("name")
    email = payload.get("email")
    phone = payload.get("phone")
    links = payload.get("links")
    company = payload.get("company")
    job_title = payload.get("job_title")
    resume_text = payload.get("resume_text", "")
    job_description = payload.get("job_description", "")
    paragraphs = payload.get("paragraphs")  # optional structured {opening, skills_fit, conclusion}

    if paragraphs is None:
        if not gemini_api_key:
            raise HTTPException(status_code=400, detail="API key not set")
        genai.configure(api_key=gemini_api_key)
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = build_cover_letter_prompt(resume_text, job_description, job_title or "", company or "")
        resp = model.generate_content(prompt)
        try:
            paragraphs = json.loads(resp.text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gemini returned non-JSON: {str(e)}")

    # Load template
    base_dir = Path(__file__).parent
    template_path = base_dir / "templates" / "cover_letter.tex"
    if not template_path.exists():
        raise HTTPException(status_code=500, detail="LaTeX template not found at backend/templates/cover_letter.tex")
    template_str = template_path.read_text(encoding="utf-8")
    tex = render_cover_letter_from_data(template_str, {
        "name": name or "",
        "email": email or "",
        "phone": phone or "",
        "links": links or "",
        "company": company or "",
        "job_title": job_title or "",
        "opening": paragraphs.get("opening", ""),
        "skills_fit": paragraphs.get("skills_fit", ""),
        "conclusion": paragraphs.get("conclusion", ""),
    })
    pdf_bytes = await compile_tex_to_pdf_bytes(tex)
    b64 = base64.b64encode(pdf_bytes).decode('ascii')
    safe_username = re.sub(r"[^A-Za-z0-9_-]+", "_", (name or "candidate")).strip("_") or "candidate"
    return {
        "status": "success",
        "filename": f"{safe_username}_cover_letter.pdf",
        "latex": tex,
        "pdf_base64": b64,
        "paragraphs": paragraphs,
    }

@app.post("/download/{format}")
async def download_resume(
    format: str,
    resume_text: str = Form(...),
    filename: str = Form(default="tailored_resume")
):
    """Download tailored resume in specified format (pdf, docx, txt)"""
    try:
        # Helper functions to offload blocking IO / CPU work
        def _create_txt_file(text: str) -> str:
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as f:
                f.write(text)
                return f.name

        def _create_docx_file(text: str) -> str:
            doc = Document()
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as f:
                doc.save(f.name)
                return f.name

        def _create_pdf_file(text: str) -> str:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as f:
                doc = SimpleDocTemplate(f.name, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                for paragraph in text.split('\n'):
                    if paragraph.strip():
                        story.append(Paragraph(paragraph.strip(), styles['Normal']))
                        story.append(Spacer(1, 12))
                doc.build(story)
                return f.name

        if format == "txt":
            temp_path = await anyio.to_thread.run_sync(_create_txt_file, resume_text)
            return FileResponse(temp_path, media_type='text/plain', filename=f"{filename}.txt")
        elif format == "docx":
            temp_path = await anyio.to_thread.run_sync(_create_docx_file, resume_text)
            return FileResponse(temp_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=f"{filename}.docx")
        elif format == "pdf":
            temp_path = await anyio.to_thread.run_sync(_create_pdf_file, resume_text)
            return FileResponse(temp_path, media_type='application/pdf', filename=f"{filename}.pdf")
        else:
            raise HTTPException(status_code=400, detail="Invalid format. Use: txt, docx, or pdf")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")

def _ensure_output_dir() -> Path:
    out = Path(__file__).parent / "output"
    out.mkdir(parents=True, exist_ok=True)
    return out

# ---------- Overleaf-style resume rendering ----------

def truncate_list(items: List, max_len: int) -> List:
    if not items:
        return []
    return items[:max_len]

def join_inline(items: List[str]) -> str:
    return ", ".join(escape_latex(str(x)) for x in items if x)

def render_resume_tex_overleaf(data: Dict) -> str:
    """Render LaTeX using Overleaf-like single-page CV layout with placeholders.
    Expected keys per spec: name, contact{github, linkedin, website, email, phone}, summary,
    experience[{title, company, date, points[]}], projects[{title, link, desc}],
    education[{date, degree, institute, gpa}], publications[{citation}],
    skills_left[list], skills_right[list], certifications[list].
    """
    # Enforce single-page by limiting counts (first pass)
    data = dict(data or {})
    data["experience"] = truncate_list(data.get("experience") or [], 3)
    for exp in data["experience"]:
        exp["points"] = truncate_list(exp.get("points") or [], 3)
    data["projects"] = truncate_list(data.get("projects") or [], 3)
    data["education"] = truncate_list(data.get("education") or [], 3)
    data["publications"] = truncate_list(data.get("publications") or [], 2)
    data["certifications"] = truncate_list(data.get("certifications") or [], 3)

    name = escape_latex(data.get("name", ""))
    contact = data.get("contact") or {}
    github = escape_latex(contact.get("github", ""))
    linkedin = escape_latex(contact.get("linkedin", ""))
    website = escape_latex(contact.get("website", ""))
    email = escape_latex(contact.get("email", ""))
    phone = escape_latex(contact.get("phone", ""))

    summary = escape_latex(data.get("summary", ""))

    # Work Experience blocks
    work_blocks = []
    for exp in data.get("experience") or []:
        title = escape_latex(exp.get("title", ""))
        company = escape_latex(exp.get("company", ""))
        date = escape_latex(exp.get("date", ""))
        points = exp.get("points") or []
        body = "\\begin{itemize}[leftmargin=*]\n" + "\n".join(
            f"  \\item {escape_latex(p)}" for p in points if p
        ) + "\n\\end{itemize}"
        work_blocks.append(
            f"\\joblong{{{title}}}{{{company}}}{{{date}}}{{%\n{body}\n}}"
        )
    work_tex = "\n\n".join(work_blocks)

    # Projects
    proj_blocks = []
    for pr in data.get("projects") or []:
        title = escape_latex(pr.get("title", ""))
        link = escape_latex(pr.get("link", ""))
        desc = escape_latex(pr.get("desc", ""))
        proj_blocks.append(f"\\project{{{title}}}{{{link}}}{{{desc}}}")
    projects_tex = "\n\n".join(proj_blocks)

    # Education two-column
    edu_lines = []
    for ed in data.get("education") or []:
        date = escape_latex(ed.get("date", ""))
        degree = escape_latex(ed.get("degree", ""))
        institute = escape_latex(ed.get("institute", ""))
        gpa = escape_latex(ed.get("gpa", ""))
        edu_lines.append(f"{date} & {degree} \\textbf{{at}} {institute} & (GPA: {gpa})\\\\")
    edu_tex = "\n".join(edu_lines)

    # Publications
    pub_items = []
    for pub in data.get("publications") or []:
        pub_items.append(f"\\item {escape_latex(pub.get('citation', ''))}")
    pubs_tex = "\\begin{itemize}[leftmargin=*]\n" + "\n".join(pub_items) + "\n\\end{itemize}" if pub_items else ""

    # Skills two-column
    sl = data.get("skills_left") or []
    sr = data.get("skills_right") or []
    skills_left = " \\tabitem ".join(escape_latex(str(x)) for x in sl) if sl else ""
    skills_right = " \\tabitem ".join(escape_latex(str(x)) for x in sr) if sr else ""

    # Certifications optional line
    certs = data.get("certifications") or []
    certs_tex = " \\tabitem ".join(escape_latex(str(x)) for x in certs) if certs else ""

    base_dir = Path(__file__).parent
    strict_tpl = base_dir / "templates" / "resume_template.tex"
    if strict_tpl.exists():
        tpl = strict_tpl.read_text(encoding="utf-8")
        repl = {
            "{{NAME}}": name,
            "{{GITHUB}}": github,
            "{{LINKEDIN}}": linkedin,
            "{{WEBSITE}}": website,
            "{{EMAIL}}": email,
            "{{PHONE}}": phone,
            "{{SUMMARY}}": summary,
            "{{WORK}}": work_tex,
            "{{PROJECTS}}": projects_tex,
            "{{EDU_ROWS}}": edu_tex,
            "{{PUBLICATIONS}}": pubs_tex,
            "{{SKILLS_LEFT}}": skills_left,
            "{{SKILLS_RIGHT}}": skills_right,
            "{{CERTIFICATIONS}}": certs_tex,
        }
        for k, v in repl.items():
            tpl = tpl.replace(k, v or "")
        return tpl
    # Fallback to simple template (resume.tex) if strict template is missing
    simple_tpl = base_dir / "templates" / "resume.tex"
    if not simple_tpl.exists():
        raise HTTPException(status_code=500, detail="LaTeX template not found (expected templates/resume_template.tex or templates/resume.tex)")
    # Convert schema to simple renderer expected keys (direct mapping)
    def split_date(d: str):
        d = d or ""
        return (d, "") if "–" not in d and "-" not in d else tuple([p.strip() for p in re.split("–|-", d, maxsplit=1)] + [""])[:2]
    simple_data = {
        "name": name,
        "email": email,
        "phone": phone,
        "links": ", ".join(filter(None, [github, linkedin, website])),
        "summary": summary,
        "skills": [*([s.strip() for s in skills_left.split("\\tabitem") if s.strip()]), *([s.strip() for s in skills_right.split("\\tabitem") if s.strip()])],
        "experience": [
            {
                "role": escape_latex(exp.get("title","")),
                "company": escape_latex(exp.get("company","")),
                "location": "",
                "start": split_date(exp.get("date",""))[0],
                "end": split_date(exp.get("date",""))[1],
                "bullets": [escape_latex(p) for p in (exp.get("points") or [])]
            } for exp in (data.get("experience") or [])
        ],
        "projects": [
            {"name": pr.get("title",""), "url": pr.get("link",""), "description": pr.get("desc",""), "bullets": []}
            for pr in (data.get("projects") or [])
        ],
        "education": [
            {"institution": escape_latex(ed.get("institute","")), "degree": escape_latex(ed.get("degree","")), "location": "", "start": "", "end": escape_latex(ed.get("date","")), "details": escape_latex(ed.get("gpa",""))}
            for ed in (data.get("education") or [])
        ],
        "certifications": [escape_latex(c) for c in (data.get("certifications") or [])]
    }
    return render_latex_from_data(simple_tpl.read_text(encoding="utf-8"), simple_data)

def predict_overflow_lines(data: Dict) -> int:
    """Very rough line estimate to keep to one page.
    Counts bullets + project lines + education lines + summary weight.
    """
    exp_lines = sum(len(e.get("points") or []) + 2 for e in (data.get("experience") or []))
    proj_lines = len(data.get("projects") or [])
    edu_lines = len(data.get("education") or [])
    sum_lines = 2 if data.get("summary") else 0
    return exp_lines + proj_lines + edu_lines + sum_lines

def prune_for_single_page(data: Dict) -> Dict:
    """Prune content if overflow predicted. Reduce bullets then trim oldest entries."""
    pruned = json.loads(json.dumps(data))  # deep copy
    # reduce bullets to 2
    for e in pruned.get("experience", [])[:]:
        pts = e.get("points") or []
        e["points"] = pts[:2]
    # keep max 2 projects one-line
    pruned["projects"] = (pruned.get("projects") or [])[:2]
    # keep max 3 experiences
    pruned["experience"] = (pruned.get("experience") or [])[:3]
    # trim education to 2
    pruned["education"] = (pruned.get("education") or [])[:2]
    return pruned

# ---------- Compilation helpers ----------
async def compile_overleaf_pdf(latex_str: str, out_pdf_name: str) -> bytes:
    out_dir = _ensure_output_dir()
    tex_path = out_dir / (Path(out_pdf_name).stem + ".tex")
    tex_path.write_text(latex_str, encoding="utf-8")
    # Use latexmk if available else pdflatex (twice)
    compiler = find_latex_compiler()
    if not compiler:
        # Remote fallback
        try:
            return compile_via_latexonline(latex_str)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"No local LaTeX compiler and remote compile failed: {str(e)}")
    if compiler == "latexmk":
        cmd = ["latexmk", "-pdf", "-interaction=nonstopmode", "-halt-on-error", tex_path.name]
        runs = 1
    else:
        cmd = ["pdflatex", "-interaction=nonstopmode", tex_path.name]
        runs = 2
    for _ in range(runs):
        proc = await asyncio.create_subprocess_exec(*cmd, cwd=str(out_dir), stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.STDOUT)
        stdout, _ = await proc.communicate()
        if proc.returncode != 0:
            # Try remote fallback if local compile fails
            try:
                return compile_via_latexonline(latex_str)
            except Exception as e:
                log = stdout.decode(errors="ignore")
                raise HTTPException(status_code=500, detail=f"Local LaTeX compile failed and remote fallback error: {str(e)}\n{log[-2000:]}")
    pdf_path = out_dir / (Path(out_pdf_name).stem + ".pdf")
    if not pdf_path.exists():
        raise HTTPException(status_code=500, detail="Compiled PDF not found")
    return pdf_path.read_bytes()

def compile_with_latexmk(latex_str: str, jobname: str = "resume") -> (bytes, int, str):
    """Compile LaTeX using latexmk with output dir 'output' and return (pdf_bytes, page_count, log_text)."""
    out_dir = _ensure_output_dir()
    tex_path = out_dir / f"{jobname}.tex"
    tex_path.write_text(latex_str, encoding="utf-8")
    # latexmk command
    cmd = [
        "latexmk", "-pdf", "-halt-on-error", "-interaction=nonstopmode",
        "-output-directory", str(out_dir), tex_path.name
    ]
    # Execute
    proc = subprocess.run(cmd, cwd=str(out_dir), capture_output=True, text=True)
    log_text = (proc.stdout or "") + "\n" + (proc.stderr or "")
    if proc.returncode != 0:
        raise HTTPException(status_code=500, detail=f"LaTeX compilation failed")
    pdf_path = out_dir / f"{jobname}.pdf"
    if not pdf_path.exists():
        raise HTTPException(status_code=500, detail="PDF not generated by LaTeX compiler.")
    # Page count from log if available
    page_count = 1
    try:
        # latexmk writes a .log next to pdf
        log_path = out_dir / f"{jobname}.log"
        if log_path.exists():
            log = log_path.read_text(errors="ignore")
            m = re.search(r"Output written on .*\((\d+) pages?\)", log)
            if m:
                page_count = int(m.group(1))
            log_text = log
    except Exception:
        pass
    return pdf_path.read_bytes(), page_count, log_text

def compile_via_latexonline(tex_source: str) -> bytes:
    """Compile LaTeX using latexonline.cc as a fallback. Returns PDF bytes."""
    base = "https://latexonline.cc/compile"
    def minify_latex(src: str) -> str:
        # remove comments (%) not escaped, collapse whitespace
        out_lines = []
        for line in (src or "").splitlines():
            i = 0
            cut = len(line)
            while i < len(line):
                if line[i] == '%':
                    # if escaped \% keep; else cut here
                    if i > 0 and line[i-1] == '\\':
                        i += 1
                        continue
                    cut = i
                    break
                i += 1
            if cut > 0:
                out_lines.append(line[:cut])
        compact = ' '.join(' '.join(out_lines).split())
        return compact
    mini = minify_latex(tex_source)
    # Try GET first (short docs recommend GET with text query)
    try:
        params = {
            "text": mini,
            "engine": "pdflatex"
        }
        r = requests.get(base, params=params, timeout=120)
        if r.status_code == 200 and 'pdf' in (r.headers.get('Content-Type', '').lower()):
            return r.content
        # Fallthrough to POST if content-type not pdf
    except Exception:
        pass
    # POST fallback to avoid URL length
    try:
        data = {"text": mini, "engine": "pdflatex", "directive": "general"}
        r = requests.post(base, data=data, timeout=180)
        if r.status_code == 200 and 'pdf' in (r.headers.get('Content-Type', '').lower()):
            return r.content
        raise RuntimeError(f"latexonline.cc returned {r.status_code} {r.headers.get('Content-Type','')}\n{(r.text or '')[:300]}")
    except Exception as e:
        raise RuntimeError(f"latexonline.cc compile failed: {e}")

# ---------- Simple PDF fallback (ReportLab) ----------
def render_simple_pdf_from_data(data: Dict) -> bytes:
    """Render a clean single-page PDF using ReportLab from structured resume data.
    This is a last-resort fallback when LaTeX compilation is unavailable or fails.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    styles = getSampleStyleSheet()
    story = []

    def add_para(text: str, style_name: str = 'Normal', space_after: int = 6):
        if text:
            story.append(Paragraph(text, styles[style_name]))
            story.append(Spacer(1, space_after))

    # Header
    name = data.get('name') or 'Candidate Name'
    contact = data.get('contact') or {}
    contact_line = " | ".join(filter(None, [
        contact.get('email') or '',
        contact.get('phone') or '',
        contact.get('website') or '',
        contact.get('github') or '',
        contact.get('linkedin') or '',
    ]))
    add_para(f"<b>{name}</b>", 'Title', 4)
    add_para(contact_line, 'Normal', 8)

    # Summary
    summary = data.get('summary') or ''
    if summary:
        add_para('<b>Summary</b>', 'Heading3', 2)
        add_para(summary, 'Normal', 8)

    # Experience
    exps = data.get('experience') or []
    if exps:
        add_para('<b>Work Experience</b>', 'Heading3', 2)
        for e in exps[:3]:
            title = (e.get('title') or e.get('role') or '')
            company = e.get('company') or ''
            date = e.get('date') or f"{e.get('start','')} – {e.get('end','')}"
            add_para(f"<b>{title}</b> — {company} <font size=9 color=grey>({date})</font>", 'Normal', 2)
            for p in (e.get('points') or e.get('bullets') or [])[:3]:
                add_para(f"• {p}", 'Normal', 1)
            story.append(Spacer(1, 6))

    # Projects
    projs = data.get('projects') or []
    if projs:
        add_para('<b>Projects</b>', 'Heading3', 2)
        for pr in projs[:3]:
            title = pr.get('title') or pr.get('name') or ''
            link = pr.get('link') or pr.get('url') or ''
            desc = pr.get('desc') or pr.get('description') or ''
            add_para(f"<b>{title}</b> {link}", 'Normal', 1)
            add_para(desc, 'Normal', 6)

    # Education
    edus = data.get('education') or []
    if edus:
        add_para('<b>Education</b>', 'Heading3', 2)
        for ed in edus[:3]:
            degree = ed.get('degree') or ''
            inst = ed.get('institute') or ed.get('institution') or ''
            gpa = ed.get('gpa') or ''
            date = ed.get('date') or ed.get('end') or ''
            add_para(f"{degree} — {inst} <font size=9 color=grey>({date})</font>", 'Normal', 1)
            if gpa:
                add_para(f"GPA: {gpa}", 'Normal', 4)

    # Skills
    sl = data.get('skills_left') or []
    sr = data.get('skills_right') or []
    skills = ", ".join([*(sl[:12]), *(sr[:12])])
    if skills:
        add_para('<b>Skills</b>', 'Heading3', 2)
        add_para(skills, 'Normal', 0)

    # Build
    doc.build(story)
    pdf = buf.getvalue()
    buf.close()
    return pdf

# ---------- Gemini JSON helpers ----------
def extract_json_object(text: str) -> Dict:
    """Try to parse JSON from raw model text. Handles code fences and prose around JSON."""
    if not text:
        raise ValueError("Empty response from model")
    # Fast path
    try:
        return json.loads(text)
    except Exception:
        pass
    # Remove code fences
    cleaned = re.sub(r"```(json|JSON)?", "", text)
    cleaned = cleaned.replace("```", "").strip()
    # Find first JSON object
    m = re.search(r"\{[\s\S]*\}", cleaned)
    if m:
        candidate = m.group(0)
        try:
            return json.loads(candidate)
        except Exception:
            pass
    raise ValueError("Could not extract valid JSON from model output")

def fallback_resume_data(resume_text: str, job_description: str) -> Dict:
    """Produce a compact, single-page-friendly resume JSON if model fails.
    Uses simple heuristics from resume_text and JD keywords.
    """
    # Very small summary (2 lines max)
    first_lines = " ".join([ln.strip() for ln in (resume_text or "").splitlines()[:4]])
    jd_words = re.findall(r"[A-Za-z]{3,}", job_description or "")
    jd_top = list(dict.fromkeys([w.lower() for w in jd_words]))[:8]
    summary = (first_lines[:240] + ("..." if len(first_lines) > 240 else "")) or "Results-oriented professional aligned to the role."

    skills_left = jd_top[:4]
    skills_right = jd_top[4:8]

    return {
        "name": "Candidate Name",
        "contact": {"github":"","linkedin":"","website":"","email":"","phone":""},
        "summary": summary,
        "experience": [
            {"title":"Relevant Experience","company":"","date":"","points":[
                "Delivered outcomes aligned to JD priorities.",
                "Collaborated across teams; improved processes.",
                "Applied tools and methods referenced in the JD."
            ]}
        ],
        "projects": [
            {"title":"Key Project","link":"","desc":"Demonstrated skills matching the JD in a real project."}
        ],
        "education": [],
        "publications": [],
        "skills_left": skills_left,
        "skills_right": skills_right,
        "certifications": []
    }

# ---------- New Endpoints ----------

@app.post("/tailor_resume_overleaf")
async def tailor_resume_overleaf(payload: Dict):
    """Call Gemini to produce structured JSON per schema, render into Overleaf template, compile PDF, return base64+LaTeX+ATS.
    Returns HTTP 500 on failure with a clear message.
    """
    try:
        global gemini_api_key
        resume_text = payload.get("resume_text", "")
        job_description = payload.get("job_description") or payload.get("jd_text", "")
        username = payload.get("username", "candidate")

        # ATS before
        ats_before = calculate_ats_score(resume_text, job_description).get("score", 0)

        # Gemini prompt per spec
        prompt = (
            "You are an expert resume writer. Using the user’s original resume and the job description, return a single-page tailored resume as JSON matching this schema: "
            "{"
            "\"name\":\"...\",\"contact\":{\"github\":\"...\",\"linkedin\":\"...\",\"website\":\"...\",\"email\":\"...\",\"phone\":\"...\"},"
            "\"summary\":\"...\",\"experience\":[{\"title\":\"...\",\"company\":\"...\",\"date\":\"...\",\"points\":[\"...\"]}],"
            "\"projects\":[{\"title\":\"...\",\"link\":\"...\",\"desc\":\"...\"}],"
            "\"education\":[{\"date\":\"...\",\"degree\":\"...\",\"institute\":\"...\",\"gpa\":\"...\"}],"
            "\"publications\":[{\"citation\":\"...\"}],"
            "\"skills_left\":[\"items\"],\"skills_right\":[\"items\"],\"certifications\":[\"...\"]"
            "}"
            " Maintain factual accuracy (no fabricated roles), prioritize JD-aligned keywords, concise bullet points (max 2 lines each), and fit one A4 page in LaTeX. Dates like 'Jan 2023 – Present'.\n\n"
            f"Original Resume:\n{resume_text}\n\nJob Description:\n{job_description}"
        )
        # Enforce Gemini usage and fail if unavailable/invalid
        if not gemini_api_key:
            raise HTTPException(status_code=400, detail="API key not set")
        if not (resume_text or "").strip():
            raise HTTPException(status_code=400, detail="resume_text is required")
        if not (job_description or "").strip():
            raise HTTPException(status_code=400, detail="job_description is required")
        try:
            genai.configure(api_key=gemini_api_key)
            model = genai.GenerativeModel('gemini-1.5-flash', generation_config={"temperature":0.4})
            resp = model.generate_content(prompt)
            data = extract_json_object(getattr(resp, 'text', '') or '')
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gemini generation failed: {e}")

        # Enforce stricter pruning if overflow predicted
        if predict_overflow_lines(data) > 28:
            data = prune_for_single_page(data)

        # Render LaTeX and compile; do not embed any debug/source blocks
        include_source_in_pdf = False
        tex = render_resume_tex_overleaf(data)
        # guard: strip accidental \newpage or \clearpage
        tex = re.sub(r"\\(newpage|clearpage)\b", "", tex)
        if not tex or not tex.strip().endswith("\\end{document}"):
            raise HTTPException(status_code=500, detail="latex_source empty after templating")
        # Diagnostic: print first 200 chars
        try:
            print("LATEX[200]: ", tex[:200])
        except Exception:
            pass

        # Fast mode: optionally skip LaTeX compilation for speed
        pdf_b64 = ""
        page_count = 0
        fast = bool(payload.get("fast") or payload.get("fast_mode"))
        if fast:
            try:
                pdf_bytes = render_simple_pdf_from_data(data)
                page_count = 1
                pdf_b64 = "data:application/pdf;base64," + base64.b64encode(pdf_bytes).decode('ascii')
            except Exception as e_fast:
                raise HTTPException(status_code=500, detail=f"Fast generation failed: {e_fast}")
        else:
            # Compile to PDF (latexmk preferred; fallback to remote)
            try:
                if shutil.which("latexmk"):
                    pdf_bytes, page_count, _ = compile_with_latexmk(tex, jobname=f"{re.sub(r'[^A-Za-z0-9_-]+','_',username)}_resume")
                    if page_count and page_count > 1:
                        data = prune_for_single_page(data)
                        tex = render_resume_tex_overleaf(data)
                        tex = re.sub(r"\\(newpage|clearpage)\\b", "", tex)
                        pdf_bytes, page_count, _ = compile_with_latexmk(tex, jobname=f"{re.sub(r'[^A-Za-z0-9_-]+','_',username)}_resume")
                else:
                    # Remote compile fallback
                    pdf_bytes = compile_via_latexonline(tex)
                    page_count = 0
                pdf_b64 = "data:application/pdf;base64," + base64.b64encode(pdf_bytes).decode('ascii')
            except Exception as e:
                # Attempt remote fallback if local path failed, then simple PDF as last resort
                try:
                    pdf_bytes = compile_via_latexonline(tex)
                    page_count = 0
                    pdf_b64 = "data:application/pdf;base64," + base64.b64encode(pdf_bytes).decode('ascii')
                except Exception:
                    # Final fallback: render simple PDF so user still gets a valid file
                    try:
                        pdf_bytes = render_simple_pdf_from_data(data)
                        page_count = 1
                        pdf_b64 = "data:application/pdf;base64," + base64.b64encode(pdf_bytes).decode('ascii')
                    except Exception as e3:
                        raise HTTPException(status_code=500, detail=f"PDF compilation failed and fallback failed: {e3}")

        # ATS after
        plain_tailored = " ".join([
            data.get("summary", ""),
            " ".join(p for e in (data.get("experience") or []) for p in (e.get("points") or [])),
            " ".join(pr.get("desc","") for pr in (data.get("projects") or [])),
            " ".join(ed.get("degree","")+" "+ed.get("institute","") for ed in (data.get("education") or [])),
            join_inline(data.get("skills_left") or []), join_inline(data.get("skills_right") or [])
        ])
        after_analysis = calculate_ats_score(plain_tailored, job_description)
        ats_after = after_analysis.get("score", 0)
        missing_after = after_analysis.get("missing_keywords", [])
        resp = {
            "status": "success",
            "filename": f"{re.sub(r'[^A-Za-z0-9_-]+','_',username)}_resume.pdf",
            "latex_source": tex,
            "pdf_base64": pdf_b64,
            "ats_before": ats_before,
            "ats_after": ats_after,
            "missing_keywords": missing_after,
            "page_count": page_count
        }
        return resp
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"tailor_resume_overleaf failed: {str(e)}")

@app.post("/generate_cover_letter_overleaf")
async def generate_cover_letter_overleaf(payload: Dict):
    global gemini_api_key
    resume_summary = payload.get("resume_summary", "")
    job_description = payload.get("job_description", "")
    company = payload.get("company", "")
    role = payload.get("role", "")
    name = payload.get("name", "")
    contact = payload.get("contact") or {}

    out = None
    if gemini_api_key:
        try:
            genai.configure(api_key=gemini_api_key)
            model = genai.GenerativeModel('gemini-1.5-flash', generation_config={"temperature":0.4})
            prompt = (
                "Generate a concise, single-page professional cover letter in 3–4 short paragraphs tailored to the JD and company. "
                "Return JSON: { 'recipient': {'company':'...','role':'...'}, 'body':['para1','para2','para3','closing'], 'signoff':'Full Name' }. "
                "Keep it ATS-friendly, factual, and action-oriented.\n\n"
                f"Resume Summary:\n{resume_summary}\n\nJob Description:\n{job_description}\n\nCompany: {company}\nRole: {role}"
            )
            resp = model.generate_content(prompt)
            out = extract_json_object(getattr(resp, 'text', '') or '')
        except Exception:
            out = None
    if not out:
        out = {"recipient": {"company": company, "role": role}, "body": [
            "I am excited to apply for the role at your company.",
            "My experience and skills align with the job description.",
            "I would welcome the opportunity to contribute and learn.",
            "Thank you for your consideration."
        ], "signoff": name or "Candidate"}

    # Assemble LaTeX cover letter using existing cover_letter.tex template
    base_dir = Path(__file__).parent
    tpl_path = base_dir / "templates" / "cover_letter.tex"
    if not tpl_path.exists():
        raise HTTPException(status_code=500, detail="cover_letter.tex not found")
    tpl = tpl_path.read_text(encoding="utf-8")

    def esc(x):
        return escape_latex(x or "")

    body = out.get("body") or []
    tex = tpl
    tex = tex.replace("{{NAME}}", esc(name))
    tex = tex.replace("{{EMAIL}}", esc(contact.get("email", "")))
    tex = tex.replace("{{PHONE}}", esc(contact.get("phone", "")))
    links = ", ".join(filter(None, [contact.get("github"), contact.get("linkedin"), contact.get("website")]))
    tex = tex.replace("{{LINKS}}", esc(links))
    tex = tex.replace("{{COMPANY}}", esc(out.get("recipient", {}).get("company", company)))
    tex = tex.replace("{{HIRING_MANAGER}}", esc("Hiring Manager"))
    tex = tex.replace("{{JOB_TITLE}}", esc(out.get("recipient", {}).get("role", role)))
    tex = tex.replace("{{OPENING}}", esc(body[0] if len(body)>0 else ""))
    tex = tex.replace("{{SKILLS_FIT}}", esc(body[1] if len(body)>1 else ""))
    tex = tex.replace("{{CONCLUSION}}", esc((body[2] if len(body)>2 else "") + (" " + body[3] if len(body)>3 else "")))

    pdf_bytes = await compile_overleaf_pdf(tex, f"{re.sub(r'[^A-Za-z0-9_-]+','_',name or 'candidate')}_cover_letter.tex")

    return {
        "status": "success",
        "filename": f"{re.sub(r'[^A-Za-z0-9_-]+','_',name or 'candidate')}_cover_letter.pdf",
        "latex": tex,
        "pdf_base64": base64.b64encode(pdf_bytes).decode('ascii')
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)